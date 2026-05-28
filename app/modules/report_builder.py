import os
from pathlib import Path
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
from app.utils import logger, BASE_DIR, CHARTS_DIR, load_config

config = load_config()
report_config = config.get("report", {})

def _format_p_value(p: float) -> str:
    """统一P值格式：小于0.001时使用科学计数法"""
    if p is None:
        return "-"
    if p < 0.001:
        return f"{p:.2e}"
    return f"{p:.4f}"

def create_word_report(
    filename: str,
    df: pd.DataFrame,
    cleaning_log: Dict[str, Any],
    analysis_results: Dict[str, Any],
    charts: Dict[str, str],
    ai_conclusion: Optional[str] = None,
    output_path: Optional[str] = None,
    issues: Optional[List[Dict[str, Any]]] = None,
    template_type: str = "academic",
    logo_path: Optional[str] = None,
    author_info: Optional[Dict[str, str]] = None,
    language: str = "zh",
    include_raw_data: bool = False,
    include_code_snippet: bool = False,
) -> str:
    """创建Word分析报告，支持多种模板和自定义选项
    
    Args:
        filename: 数据文件名
        df: 清洗后的数据框
        cleaning_log: 清洗日志
        analysis_results: 分析结果字典
        charts: 图表路径字典
        ai_conclusion: AI生成的结论
        output_path: 输出路径（可选）
        issues: 问题诊断列表
        template_type: 报告模板类型（academic/simple/clinical）
        logo_path: Logo图片路径
        author_info: 作者信息字典（author, institution, project_id）
        language: 报告语言（zh/en）
        include_raw_data: 是否包含原始数据预览
        include_code_snippet: 是否包含Python代码片段
    
    Returns:
        报告文件路径
    """
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = report_config.get("font", {}).get("english", "Times New Roman")
    font.size = Pt(report_config.get("font_size", 12))

    # 多语言支持
    lang_map = {
        "title": {
            "zh": "生物医药实验数据分析报告",
            "en": "Biomedical Experimental Data Analysis Report",
        },
        "data_file": {"zh": "数据文件", "en": "Data File"},
        "data_rows": {"zh": "数据行数", "en": "Data Rows"},
        "data_cols": {"zh": "列数", "en": "Columns"},
        "group_var": {"zh": "分组变量", "en": "Grouping Variable"},
        "group_size": {"zh": "各组样本量", "en": "Group Sample Sizes"},
        "sig_level": {"zh": "显著性水平", "en": "Significance Level"},
        "purpose": {"zh": "研究目的与假设", "en": "Research Purpose & Hypotheses"},
        "purpose_text": {
            "zh": "本研究旨在通过统计分析方法，探索实验数据中各变量之间的关系及组间差异。",
            "en": "This study aims to explore the relationships between variables and group differences in experimental data through statistical analysis methods.",
        },
        "hypothesis": {
            "zh": "核心假设：不同{group}组间在关键指标上存在显著差异。",
            "en": "Core hypothesis: There are significant differences in key indicators among different {group} groups.",
        },
        "note": {
            "zh": "注：本报告基于观察性数据分析，相关性不等于因果性。",
            "en": "Note: This report is based on observational data analysis; correlation does not imply causation.",
        },
        "cleaning": {"zh": "一、数据清洗摘要", "en": "I. Data Cleaning Summary"},
        "original_data": {"zh": "原始数据", "en": "Original Data"},
        "cleaned_data": {"zh": "清洗后数据", "en": "Cleaned Data"},
        "cleaning_steps": {"zh": "清洗步骤", "en": "Cleaning Steps"},
        "methods": {"zh": "二、分析方法选择", "en": "II. Analysis Methods Selection"},
        "methods_text": {
            "zh": "平台已根据数据特征自动选择了以下分析方法：",
            "en": "The platform has automatically selected the following analysis methods based on data characteristics:",
        },
        "desc_stats": {"zh": "三、描述性统计", "en": "III. Descriptive Statistics"},
        "stat_results": {"zh": "四、统计分析结果", "en": "IV. Statistical Analysis Results"},
        "charts": {"zh": "五、图表", "en": "V. Charts & Figures"},
        "ai_conclusion": {"zh": "六、AI分析结论", "en": "VI. AI Analysis Conclusion"},
        "diagnosis": {"zh": "六、数据质量与分析问题诊断及建议", "en": "VI. Data Quality & Analysis Issues Diagnosis"},
        "limitations": {"zh": "七、研究局限性与建议", "en": "VII. Research Limitations & Recommendations"},
        "limitations_title": {"zh": "局限性", "en": "Limitations"},
        "suggestions_title": {"zh": "后续建议", "en": "Recommendations"},
        "appendix": {"zh": "八、附录", "en": "VIII. Appendix"},
        "raw_data": {"zh": "原始数据预览（前10行）", "en": "Raw Data Preview (First 10 Rows)"},
        "code_snippet": {"zh": "Python分析代码片段", "en": "Python Analysis Code Snippet"},
    }

    def t(key: str, **kwargs) -> str:
        """翻译函数"""
        text = lang_map.get(key, {}).get(language, lang_map.get(key, {}).get("zh", key))
        if kwargs:
            for k, v in kwargs.items():
                text = text.replace(f"{{{k}}}", str(v))
        return text

    # 根据模板类型调整格式
    if template_type == "simple":
        font.size = Pt(11)
    elif template_type == "clinical":
        font.name = "Arial"
        font.size = Pt(11)

    # 封面页（学术型和临床型）
    if template_type in ["academic", "clinical"]:
        # 添加Logo
        if logo_path and os.path.exists(logo_path):
            try:
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_para.add_run().add_picture(logo_path, width=Inches(2.0))
                doc.add_paragraph("")
            except Exception as e:
                logger.warning(f"Failed to add logo: {e}")

        title = doc.add_heading(t("title"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.font.size = Pt(18)

        # 作者信息
        if author_info:
            doc.add_paragraph("")
            if author_info.get("author"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(author_info["author"]).bold = True
            
            if author_info.get("institution"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(author_info["institution"])
            
            if author_info.get("project_id"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(f"{t('data_file')}: {author_info['project_id']}")
            
            doc.add_paragraph("")
    else:
        # 简洁型：直接显示标题
        title = doc.add_heading(t("title"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

    # 基础信息
    doc.add_paragraph(f"{t('data_file')}: {filename}")
    doc.add_paragraph(f"{t('data_rows')}: {len(df)}, {t('data_cols')}: {len(df.columns)}")
    
    # 样本分组信息
    group_col = None
    for col in df.columns:
        if df[col].dtype == 'object' and df[col].nunique() <= 10:
            group_col = col
            break
    
    if group_col:
        doc.add_paragraph(f"{t('group_var')}: {group_col}")
        group_counts = df[group_col].value_counts()
        group_info = ", ".join([f"{name}: {count}例" for name, count in group_counts.items()])
        doc.add_paragraph(f"{t('group_size')}: {group_info}")
    
    doc.add_paragraph(f"{t('sig_level')}: α = 0.05")
    doc.add_paragraph("")
    
    # 研究目的与假设
    doc.add_heading(t("purpose"), level=1)
    doc.add_paragraph(t("purpose_text"))
    if group_col:
        doc.add_paragraph(t("hypothesis", group=group_col))
    doc.add_paragraph(t("note"))
    doc.add_paragraph("")

    doc.add_heading(t("cleaning"), level=1)
    doc.add_paragraph(f"{t('original_data')}: {cleaning_log.get('original_rows', 'N/A')} 行 × {cleaning_log.get('original_cols', 'N/A')} 列")
    doc.add_paragraph(f"{t('cleaned_data')}: {cleaning_log.get('final_rows', 'N/A')} 行 × {cleaning_log.get('final_cols', 'N/A')} 列")

    if cleaning_log.get("steps"):
        doc.add_paragraph(t("cleaning_steps"))
        
        # 多语言步骤名称映射
        step_name_map = {
            "zh": {
                "remove_empty_columns": "删除空列",
                "remove_empty_rows": "删除空行",
                "remove_duplicates": "删除重复行",
                "fill_missing": "填充缺失值",
                "fill_missing_categorical": "填充分类缺失值",
                "missing_summary": "缺失值汇总",
                "convert_types": "数据类型转换",
                "remove_outliers": "移除异常值",
                "drop_missing_rows": "删除含缺失值的行",
                "fill_missing_knn": "KNN插补缺失值",
                "winsorize_outliers": "异常值缩尾处理",
                "mark_outliers_as_na": "异常值标记为缺失",
                "fill_missing_after_outlier_marking": "异常值标记后重新填充",
                "outlier_detection": "异常值检测",
            },
            "en": {
                "remove_empty_columns": "Remove Empty Columns",
                "remove_empty_rows": "Remove Empty Rows",
                "remove_duplicates": "Remove Duplicates",
                "fill_missing": "Fill Missing Values",
                "fill_missing_categorical": "Fill Categorical Missing Values",
                "missing_summary": "Missing Values Summary",
                "convert_types": "Convert Data Types",
                "remove_outliers": "Remove Outliers",
                "drop_missing_rows": "Drop Rows with Missing Values",
                "fill_missing_knn": "KNN Imputation",
                "winsorize_outliers": "Winsorize Outliers",
                "mark_outliers_as_na": "Mark Outliers as NA",
                "fill_missing_after_outlier_marking": "Refill Missing Values After Outlier Marking",
                "outlier_detection": "Outlier Detection",
            },
        }
        
        lang_steps = step_name_map.get(language, step_name_map["zh"])
        
        for step in cleaning_log["steps"]:
            step_name = step.get('step', '')
            display_name = lang_steps.get(step_name, step_name)
            step_desc = ""
            if step_name == "remove_empty_columns":
                step_desc = f"删除了 {step.get('count', 0)} 个空列" if language == "zh" else f"Removed {step.get('count', 0)} empty columns"
            elif step_name == "remove_empty_rows":
                step_desc = f"删除了 {step.get('removed_count', 0)} 个空行" if language == "zh" else f"Removed {step.get('removed_count', 0)} empty rows"
            elif step_name == "remove_duplicates":
                step_desc = f"删除了 {step.get('removed_count', 0)} 个重复行" if language == "zh" else f"Removed {step.get('removed_count', 0)} duplicate rows"
            elif step_name == "fill_missing":
                step_desc = f"列 '{step.get('column', '')}' 使用 {step.get('strategy', '')} 填充了 {step.get('filled_count', 0)} 个缺失值" if language == "zh" else f"Column '{step.get('column', '')}' filled {step.get('filled_count', 0)} missing values using {step.get('strategy', '')}"
            elif step_name == "fill_missing_categorical":
                step_desc = f"列 '{step.get('column', '')}' 使用众数填充了 {step.get('filled_count', 0)} 个缺失值" if language == "zh" else f"Column '{step.get('column', '')}' filled {step.get('filled_count', 0)} missing values using mode"
            elif step_name == "missing_summary":
                step_desc = f"清洗前 {step.get('before', 0)} 个 → 清洗后 {step.get('after', 0)} 个" if language == "zh" else f"Before: {step.get('before', 0)} → After: {step.get('after', 0)}"
            elif step_name == "convert_types":
                step_desc = f"转换了 {step.get('count', 0)} 个列的数据类型" if language == "zh" else f"Converted {step.get('count', 0)} column data types"
            elif step_name == "remove_outliers":
                step_desc = f"移除了 {step.get('count', 0)} 个异常值" if language == "zh" else f"Removed {step.get('count', 0)} outlier values"
            elif step_name == "outlier_detection":
                step_desc = (
                    f"检测方法: {step.get('method', '')}，"
                    f"发现 {step.get('total_outliers', 0)} 个异常值（{step.get('outlier_percentage', '0%')}），"
                    f"涉及变量: {', '.join(step.get('outlier_columns', []))}"
                ) if language == "zh" else (
                    f"Method: {step.get('method', '')}, "
                    f"Found {step.get('total_outliers', 0)} outliers ({step.get('outlier_percentage', '0%')}), "
                    f"Variables: {', '.join(step.get('outlier_columns', []))}"
                )
            else:
                step_desc = str({k: v for k, v in step.items() if k != "step"})
            
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{display_name}: {step_desc}")

    # 分析方法选择说明
    doc.add_heading(t("methods"), level=1)
    methods_used = analysis_results.get("methods_used", [])
    if methods_used:
        doc.add_paragraph(t("methods_text"))
        for idx, method in enumerate(methods_used, 1):
            if isinstance(method, dict):
                p = doc.add_paragraph(style="List Number")
                p.add_run(f"{method['name']}").bold = True
                doc.add_paragraph(method.get("reason", ""), style="List Bullet")
            else:
                doc.add_paragraph(str(method), style="List Bullet")
    else:
        doc.add_paragraph(t("methods_text"))
        default_methods = {
            "zh": [
                "描述性统计（均值、标准差、中位数等）",
                "正态性检验（Shapiro-Wilk）",
                "组间差异检验（T检验/方差分析）",
                "相关性分析",
                "回归分析",
            ],
            "en": [
                "Descriptive Statistics (mean, std, median, etc.)",
                "Normality Test (Shapiro-Wilk)",
                "Group Difference Test (t-test/ANOVA)",
                "Correlation Analysis",
                "Regression Analysis",
            ],
        }
        for m in default_methods.get(language, default_methods["zh"]):
            doc.add_paragraph(m, style="List Bullet")
    doc.add_paragraph("")

    doc.add_heading(t("desc_stats"), level=1)
    desc_stats = analysis_results.get("descriptive_statistics") or analysis_results.get("descriptive", {})
    if desc_stats and "error" not in desc_stats:
        table = doc.add_table(rows=1, cols=7)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers_zh = ["变量", "均值", "标准差", "中位数", "最小值", "最大值", "缺失值"]
        headers_en = ["Variable", "Mean", "Std Dev", "Median", "Min", "Max", "Missing"]
        headers = headers_en if language == "en" else headers_zh
        
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            for paragraph in table.rows[0].cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for col_name, stats_data in desc_stats.items():
            if isinstance(stats_data, dict) and "mean" in stats_data:
                row = table.add_row()
                row.cells[0].text = str(col_name)
                row.cells[1].text = f"{stats_data.get('mean', 0):.3f}"
                row.cells[2].text = f"{stats_data.get('std', 0):.3f}"
                row.cells[3].text = f"{stats_data.get('median', 0):.3f}"
                row.cells[4].text = f"{stats_data.get('min', 0):.3f}"
                row.cells[5].text = f"{stats_data.get('max', 0):.3f}"
                row.cells[6].text = f"{stats_data.get('missing', 0)}"

    doc.add_heading(t("stat_results"), level=1)
    
    skip_keys = ["descriptive_statistics", "descriptive", "charts", "interactive_charts", "filename", "cleaning_log", "data_characteristics", "analysis_log", "methods_used"]
    
    title_map = {
        "zh": {
            "normality": "正态性检验",
            "normality_tests": "正态性检验",
            "t_test": "T检验",
            "t_tests": "T检验",
            "anova": "方差分析",
            "anova_tests": "方差分析",
            "correlation": "相关性分析",
            "correlation_analysis": "相关性分析",
            "regression": "回归分析",
            "regression_analysis": "回归分析",
        },
        "en": {
            "normality": "Normality Test",
            "normality_tests": "Normality Test",
            "t_test": "t-Test",
            "t_tests": "t-Tests",
            "anova": "ANOVA",
            "anova_tests": "ANOVA Tests",
            "correlation": "Correlation Analysis",
            "correlation_analysis": "Correlation Analysis",
            "regression": "Regression Analysis",
            "regression_analysis": "Regression Analysis",
        },
    }
    
    lang_titles = title_map.get(language, title_map["zh"])
    
    field_map_zh = {
        "test": "检验方法",
        "is_normal": "是否正态",
        "statistic": "统计量",
        "p_value": "P值",
        "suggestion": "建议",
        "significant": "是否显著",
        "effect_size": "效应量",
        "cohens_d": "Cohen's d",
        "eta_squared": "η²",
        "n_groups": "组数",
        "group_means": "各组均值",
        "correlation_matrix": "相关系数矩阵",
        "columns": "变量",
        "significant_pairs": "显著相关对",
        "var1": "变量1",
        "var2": "变量2",
        "correlation": "相关系数",
        "r_squared": "R²",
        "adjusted_r_squared": "调整后R²",
        "f_statistic": "F统计量",
        "f_p_value": "F检验P值",
        "rmse": "均方根误差",
        "n_samples": "样本量",
        "coefficients": "回归系数",
        "post_hoc": "事后检验",
        "method": "方法",
        "comparisons": "比较结果",
        "group1": "组1",
        "group2": "组2",
        "mean_diff": "均值差异",
        "normality": "正态性",
        "equal_variance": "方差齐性",
        "levene_p": "Levene检验P值",
        "group1_mean": "组1均值",
        "group2_mean": "组2均值",
        "group1_n": "组1样本数",
        "group2_n": "组2样本数",
        "n_pairs": "配对数",
        "group1_median": "组1中位数",
        "group2_median": "组2中位数",
        "p_value_matrix": "P值矩阵",
        "alpha": "正则化参数α",
        "l1_ratio": "L1比例",
        "n_features_selected": "选中特征数",
        "degree": "多项式次数",
        "accuracy": "准确率",
        "odds_ratio": "优势比(OR值)",
        "n_positive": "正类样本数",
        "n_negative": "负类样本数",
        "selected_features": "选中特征",
        "p_values": "P值",
        "direction": "方向",
        "model": "模型",
    }
    
    field_map_en = {
        "test": "Test Method",
        "is_normal": "Is Normal",
        "statistic": "Statistic",
        "p_value": "P-value",
        "suggestion": "Suggestion",
        "significant": "Significant",
        "effect_size": "Effect Size",
        "cohens_d": "Cohen's d",
        "eta_squared": "η²",
        "n_groups": "Number of Groups",
        "group_means": "Group Means",
        "correlation_matrix": "Correlation Matrix",
        "columns": "Variables",
        "significant_pairs": "Significant Pairs",
        "var1": "Variable 1",
        "var2": "Variable 2",
        "correlation": "Correlation",
        "r_squared": "R²",
        "adjusted_r_squared": "Adjusted R²",
        "f_statistic": "F-statistic",
        "f_p_value": "F-test P-value",
        "rmse": "RMSE",
        "n_samples": "Sample Size",
        "coefficients": "Coefficients",
        "post_hoc": "Post-hoc Test",
        "method": "Method",
        "comparisons": "Comparisons",
        "group1": "Group 1",
        "group2": "Group 2",
        "mean_diff": "Mean Difference",
        "normality": "Normality",
        "equal_variance": "Equal Variance",
        "levene_p": "Levene's P-value",
        "group1_mean": "Group 1 Mean",
        "group2_mean": "Group 2 Mean",
        "group1_n": "Group 1 N",
        "group2_n": "Group 2 N",
        "n_pairs": "Number of Pairs",
        "group1_median": "Group 1 Median",
        "group2_median": "Group 2 Median",
        "p_value_matrix": "P-value Matrix",
        "alpha": "Alpha",
        "l1_ratio": "L1 Ratio",
        "n_features_selected": "Features Selected",
        "degree": "Degree",
        "accuracy": "Accuracy",
        "odds_ratio": "Odds Ratio (OR)",
        "n_positive": "Positive Cases",
        "n_negative": "Negative Cases",
        "selected_features": "Selected Features",
        "p_values": "P-values",
        "direction": "Direction",
        "model": "Model",
    }
    
    field_map = field_map_en if language == "en" else field_map_zh
    
    has_results = False
    for test_name, result in analysis_results.items():
        if test_name in skip_keys:
            continue
        if isinstance(result, dict) and result and "error" not in result:
            has_results = True
            title = lang_titles.get(test_name, test_name.replace("_", " ").title())
            doc.add_heading(title, level=2)
            
            if test_name in ["correlation_analysis", "correlation"]:
                _add_correlation_to_doc(doc, result, field_map)
            elif test_name in ["normality_tests", "normality"]:
                _add_normality_to_doc(doc, result, field_map)
            elif test_name in ["t_tests", "t_test"]:
                _add_ttest_to_doc(doc, result, field_map)
            elif test_name in ["anova_tests", "anova"]:
                _add_anova_to_doc(doc, result, field_map)
            elif test_name in ["regression_analysis", "regression"]:
                _add_regression_to_doc(doc, result, field_map)
            else:
                _add_dict_to_doc(doc, result, field_map=field_map)
    
    if not has_results:
        no_results_text = "未选择额外的统计分析选项，或分析结果为空。" if language == "zh" else "No additional statistical analysis options selected, or analysis results are empty."
        doc.add_paragraph(no_results_text)

    doc.add_heading(t("charts"), level=1)
    chart_captions = charts.get("chart_captions", {})
    chart_num = 0
    for chart_name, chart_path in charts.items():
        if chart_name == "chart_captions":
            continue
        if not Path(chart_path).is_absolute():
            chart_path = str(CHARTS_DIR / chart_path)
        
        if os.path.exists(chart_path):
            chart_num += 1
            caption_info = chart_captions.get(chart_name, {})
            caption_title = caption_info.get("title", f"图{chart_num}")
            caption_note = caption_info.get("note", "")
            
            # 图名包含图表实际名称
            chart_display_name = chart_name.replace("_", " ").title()
            full_title = f"{caption_title} {chart_display_name}"
            doc.add_heading(full_title, level=3)
            try:
                doc.add_picture(chart_path, width=Inches(5.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logger.error(f"Failed to add chart {chart_name}: {e}")
                fail_text = f"[图表加载失败: {chart_name}]" if language == "zh" else f"[Chart loading failed: {chart_name}]"
                doc.add_paragraph(fail_text)
            
            # 图下方文本标注
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"图{chart_num} {chart_display_name}")
            run.bold = True
            run.font.size = Pt(10)
            
            if caption_note:
                p_note = doc.add_paragraph()
                p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_note = p_note.add_run(caption_note)
                run_note.italic = True
                run_note.font.size = Pt(10)
        else:
            logger.warning(f"Chart file not found: {chart_path}")
            not_found_text = f"[图表文件不存在: {chart_name}]" if language == "zh" else f"[Chart file not found: {chart_name}]"
            doc.add_paragraph(not_found_text)

    if ai_conclusion:
        doc.add_heading(t("ai_conclusion"), level=1)
        doc.add_paragraph(ai_conclusion)

    if issues:
        doc.add_heading(t("diagnosis"), level=1)
        
        severity_order = {"高": 0, "中等": 1, "低": 2}
        sorted_issues = sorted(issues, key=lambda x: severity_order.get(x.get("severity", "低"), 3))
        
        for issue in sorted_issues:
            severity_icon = {"高": "🔴", "中等": "🟡", "低": "🟢"}.get(issue.get("severity", "低"), "⚪")
            p = doc.add_paragraph()
            p.add_run(f"{severity_icon} [{issue.get('severity', '')}] {issue.get('category', '')}").bold = True
            
            problem_label = "问题" if language == "zh" else "Problem"
            suggestion_label = "建议" if language == "zh" else "Suggestion"
            details_label = "详情" if language == "zh" else "Details"
            
            p2 = doc.add_paragraph()
            p2.add_run(f"{problem_label}: ").bold = True
            p2.add_run(issue.get("problem", ""))
            
            p3 = doc.add_paragraph()
            p3.add_run(f"{suggestion_label}: ").bold = True
            p3.add_run(issue.get("suggestion", ""))
            
            if issue.get("details"):
                p4 = doc.add_paragraph()
                p4.add_run(f"{details_label}: ").bold = True
                p4.add_run(issue.get("details", ""))
            
            doc.add_paragraph("")

    doc.add_heading(t("limitations"), level=1)
    doc.add_heading(t("limitations_title"), level=2)
    limitations_zh = []
    limitations_en = []
    if len(df) < 30:
        limitations_zh.append("样本量较小（<30），可能影响统计检验效能，建议谨慎解读结果。")
        limitations_en.append("Small sample size (<30) may affect statistical power; interpret results with caution.")
    if len(df) < 100:
        limitations_zh.append("样本量有限，可能存在抽样误差，建议扩大样本量以提高结果可靠性。")
        limitations_en.append("Limited sample size may introduce sampling error; consider expanding sample size for more reliable results.")
    limitations_zh.append("本报告基于观察性数据分析，相关性不等于因果性，结论仅供参考。")
    limitations_en.append("This report is based on observational data analysis; correlation does not imply causation.")
    limitations_zh.append("未考虑潜在的混杂因素，可能影响结果的准确性。")
    limitations_en.append("Potential confounding factors were not considered, which may affect result accuracy.")
    
    limitations = limitations_en if language == "en" else limitations_zh
    for lim in limitations:
        doc.add_paragraph(lim, style="List Bullet")
    
    doc.add_heading(t("suggestions_title"), level=2)
    suggestions_zh = [
        "建议扩大样本量，提高统计检验效能。",
        "建议进行独立样本验证，确认结果的稳定性。",
        "建议结合临床/实验背景，深入分析显著结果的生物学意义。",
    ]
    suggestions_en = [
        "Expand sample size to improve statistical power.",
        "Conduct independent sample validation to confirm result stability.",
        "Integrate clinical/experimental context to deeply analyze the biological significance of significant results.",
    ]
    suggestions = suggestions_en if language == "en" else suggestions_zh
    for sug in suggestions:
        doc.add_paragraph(sug, style="List Bullet")

    if include_raw_data or include_code_snippet:
        doc.add_heading(t("appendix"), level=1)
    
    if include_raw_data:
        doc.add_heading(t("raw_data"), level=2)
        preview_df = df.head(10)
        table = doc.add_table(rows=len(preview_df) + 1, cols=len(preview_df.columns))
        table.style = "Light Grid Accent 1"
        
        for i, col in enumerate(preview_df.columns):
            table.rows[0].cells[i].text = str(col)
            for paragraph in table.rows[0].cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        for i, (_, row_data) in enumerate(preview_df.iterrows()):
            for j, val in enumerate(row_data):
                table.rows[i + 1].cells[j].text = str(val)
    
    if include_code_snippet:
        doc.add_heading(t("code_snippet"), level=2)
        code_text = """# Python 分析代码示例
import pandas as pd
from app.modules.cleaner import clean_data
from app.modules.stat_tools import descriptive_statistics, correlation_analysis
from app.modules.report_builder import create_word_report

# 加载数据
df = pd.read_csv('data.csv')

# 数据清洗
cleaned_df, cleaning_log = clean_data(df, missing_strategy='median', outlier_strategy='detect_only')

# 描述性统计
desc_stats = descriptive_statistics(cleaned_df)

# 相关性分析
corr_results = correlation_analysis(cleaned_df)

# 生成报告
report_path = create_word_report(
    filename='data.csv',
    df=cleaned_df,
    cleaning_log=cleaning_log,
    analysis_results={'descriptive_statistics': desc_stats, 'correlation_analysis': corr_results},
    charts={},
    template_type='academic',
    language='zh'
)"""
        p = doc.add_paragraph()
        p.add_run(code_text).font.name = "Courier New"

    if output_path is None:
        output_path = BASE_DIR / f"report_{filename.replace('.', '_')}.docx"

    doc.save(str(output_path))
    logger.info(f"Word report saved: {output_path}")
    return str(output_path)

def _add_correlation_to_doc(doc: Document, data: Dict[str, Any], field_map: Dict[str, str] = None):
    """将相关性分析结果以表格形式添加到文档"""
    if field_map is None:
        field_map = {}
    
    # 添加方法说明
    method = data.get("method", "pearson")
    p = doc.add_paragraph()
    p.add_run(f"方法: ").bold = True
    p.add_run(method)
    
    # 添加变量列表
    columns = data.get("columns", [])
    if columns:
        p = doc.add_paragraph()
        p.add_run("变量: ").bold = True
        p.add_run(", ".join(columns))
    
    # 相关系数矩阵表格
    corr_matrix = data.get("correlation_matrix")
    if corr_matrix and columns:
        doc.add_heading("相关系数矩阵", level=3)
        n = len(columns)
        table = doc.add_table(rows=n + 1, cols=n + 1)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        header_row = table.rows[0]
        header_row.cells[0].text = "变量"
        for j, col in enumerate(columns):
            header_row.cells[j + 1].text = col
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # 数据行
        for i, col in enumerate(columns):
            row = table.rows[i + 1]
            row.cells[0].text = col
            for j in range(n):
                val = corr_matrix[i][j]
                row.cells[j + 1].text = f"{val:.4f}" if val is not None else "-"
    
    # P值矩阵表格
    p_matrix = data.get("p_value_matrix")
    if p_matrix and columns:
        doc.add_heading("P值矩阵", level=3)
        n = len(columns)
        table = doc.add_table(rows=n + 1, cols=n + 1)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        header_row = table.rows[0]
        header_row.cells[0].text = "变量"
        for j, col in enumerate(columns):
            header_row.cells[j + 1].text = col
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # 数据行
        for i, col in enumerate(columns):
            row = table.rows[i + 1]
            row.cells[0].text = col
            for j in range(n):
                val = p_matrix[i][j]
                row.cells[j + 1].text = _format_p_value(val) if val is not None else "-"
    
    # 显著相关对（精简输出：仅显示P < 0.05的显著相关对）
    significant_pairs = data.get("significant_pairs", [])
    if significant_pairs:
        sig_pairs = [p for p in significant_pairs if p.get("significant")]
        non_sig_count = len(significant_pairs) - len(sig_pairs)
        
        doc.add_heading("显著相关对 (P < 0.05)", level=3)
        
        if sig_pairs:
            # 添加显著性说明
            p = doc.add_paragraph()
            p.add_run(f"共检测到 {len(significant_pairs)} 对变量相关性，其中 {len(sig_pairs)} 对显著（P < 0.05），{non_sig_count} 对不显著。").font.size = Pt(10)
            
            # 显著相关对表格
            table = doc.add_table(rows=len(sig_pairs) + 1, cols=5)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            headers = ["变量1", "变量2", "相关系数", "P值", "相关强度"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                for paragraph in table.rows[0].cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            for idx, pair in enumerate(sig_pairs):
                row = table.rows[idx + 1]
                row.cells[0].text = pair.get("var1", "")
                row.cells[1].text = pair.get("var2", "")
                corr = pair.get("correlation", 0)
                row.cells[2].text = f"{corr:.4f}"
                p_val = pair.get("p_value")
                row.cells[3].text = _format_p_value(p_val) if p_val is not None else "-"
                
                # 相关强度解读并添加显著性标记
                abs_corr = abs(corr)
                if abs_corr < 0.3:
                    strength = "弱相关*"
                elif abs_corr < 0.5:
                    strength = "中等相关*"
                elif abs_corr < 0.7:
                    strength = "较强相关*"
                else:
                    strength = "强相关*"
                direction = "正" if corr > 0 else "负"
                row.cells[4].text = f"{direction}{strength}"
        else:
            doc.add_paragraph(f"共检测到 {len(significant_pairs)} 对变量相关性，但无显著相关对（P < 0.05）。")
            if non_sig_count > 0:
                doc.add_paragraph(f"其余 {non_sig_count} 对相关系数不显著。")
    
    doc.add_paragraph("相关强度标准：|r| < 0.3 弱相关，0.3-0.5 中等相关，0.5-0.7 较强相关，≥ 0.7 强相关；* 表示 P < 0.05")
    doc.add_paragraph("注：相关性不等于因果性，显著相关不代表变量间存在因果关系")

def _add_normality_to_doc(doc: Document, data: Dict[str, Any], field_map: Dict[str, str] = None):
    """将正态性检验结果以表格形式添加到文档"""
    if not data:
        doc.add_paragraph("无正态性检验结果")
        return
    
    table = doc.add_table(rows=len(data) + 1, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["变量", "检验方法", "统计量", "P值", "是否正态"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for paragraph in table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for idx, (col, result) in enumerate(data.items()):
        row = table.rows[idx + 1]
        row.cells[0].text = col
        row.cells[1].text = result.get("test", "-")
        row.cells[2].text = f"{result.get('statistic', 0):.4f}" if result.get('statistic') is not None else "-"
        row.cells[3].text = f"{result.get('p_value', 0):.6f}" if result.get('p_value') is not None else "-"
        is_normal = "是" if result.get("is_normal") else "否"
        row.cells[4].text = is_normal

def _add_ttest_to_doc(doc: Document, data: Dict[str, Any], field_map: Dict[str, str] = None):
    """将T检验结果以表格形式添加到文档"""
    if not data:
        doc.add_paragraph("无T检验结果")
        return
    
    table = doc.add_table(rows=len(data) + 1, cols=7)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["比较组", "统计量", "P值", "是否显著", "Cohen's d", "效应量解读", "检验方法"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for paragraph in table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for idx, (name, result) in enumerate(data.items()):
        row = table.rows[idx + 1]
        row.cells[0].text = name
        row.cells[1].text = f"{result.get('statistic', 0):.4f}" if result.get('statistic') is not None else "-"
        p_value = result.get("p_value")
        row.cells[2].text = _format_p_value(p_value) if p_value is not None else "-"
        sig = "是*" if result.get("significant") else "否"
        row.cells[3].text = sig
        effect = result.get("effect_size", {})
        cohens_d = effect.get("cohens_d")
        row.cells[4].text = f"{cohens_d:.4f}" if cohens_d is not None else "-"
        
        # 效应量解读
        if cohens_d is not None:
            abs_d = abs(cohens_d)
            if abs_d < 0.2:
                effect_interp = "极小效应"
            elif abs_d < 0.5:
                effect_interp = "小效应"
            elif abs_d < 0.8:
                effect_interp = "中等效应"
            else:
                effect_interp = "大效应"
            row.cells[5].text = effect_interp
        else:
            row.cells[5].text = "-"
        
        row.cells[6].text = result.get("test", "-")
    
    doc.add_paragraph("* P < 0.05 表示差异显著")
    doc.add_paragraph("效应量标准：Cohen's d < 0.2 极小，0.2-0.5 小效应，0.5-0.8 中等效应，≥ 0.8 大效应")

def _add_anova_to_doc(doc: Document, data: Dict[str, Any], field_map: Dict[str, str] = None):
    """将方差分析结果以表格形式添加到文档"""
    if not data:
        doc.add_paragraph("无方差分析结果")
        return
    
    # 检验方法说明
    test_name = data.get("test", "One-way ANOVA")
    p = doc.add_paragraph()
    p.add_run("检验方法: ").bold = True
    p.add_run(test_name)
    
    # 方差齐性检验
    levene_p = data.get("levene_p")
    equal_var = data.get("equal_variance")
    if levene_p is not None:
        p = doc.add_paragraph()
        p.add_run("方差齐性检验 (Levene): ").bold = True
        p.add_run(f"P = {levene_p:.6f}")
        if equal_var:
            p.add_run("，方差齐性假设成立")
        else:
            p.add_run("，方差不齐，采用非参数检验")
    
    # 正态性检验说明
    all_normal = data.get("all_normal")
    if all_normal is not None:
        p = doc.add_paragraph()
        p.add_run("正态性检验: ").bold = True
        p.add_run("所有组数据服从正态分布" if all_normal else "部分组数据不服从正态分布")
    
    # 方法选择说明
    interpretation = data.get("interpretation")
    if interpretation:
        doc.add_paragraph(interpretation)
    
    # 主检验结果表格
    table = doc.add_table(rows=2, cols=6)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["检验方法", "统计量", "P值", "是否显著", "η²", "组数"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for paragraph in table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    row = table.rows[1]
    row.cells[0].text = test_name
    stat = data.get("statistic")
    if stat is not None:
        row.cells[1].text = f"{stat:.4f}"
    else:
        row.cells[1].text = "-"
    p_value = data.get("p_value")
    if p_value is not None:
        row.cells[2].text = _format_p_value(p_value)
    else:
        row.cells[2].text = "-"
    sig = "是*" if data.get("significant") else "否"
    row.cells[3].text = sig
    effect = data.get("effect_size", {})
    eta_sq = effect.get("eta_squared")
    if eta_sq is not None:
        row.cells[4].text = f"{eta_sq:.4f}"
    else:
        row.cells[4].text = "-"
    row.cells[5].text = str(data.get("n_groups", "-"))
    
    # 效应量解读
    effect_interp = data.get("effect_interpretation")
    if effect_interp:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("效应量解读: ").bold = True
        p.add_run(effect_interp)
    
    # 异常统计值说明
    stat_note = data.get("statistic_note")
    if stat_note:
        doc.add_paragraph(stat_note)
    p_value_note = data.get("p_value_note")
    if p_value_note:
        doc.add_paragraph(p_value_note)
    
    # 事后检验
    post_hoc = data.get("post_hoc")
    if post_hoc:
        doc.add_heading("事后检验", level=3)
        method = post_hoc.get("method", "事后检验")
        doc.add_paragraph(f"方法: {method}")
        
        comparisons = post_hoc.get("comparisons", [])
        if comparisons:
            table = doc.add_table(rows=len(comparisons) + 1, cols=5)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            headers = ["比较组", "均值差异", "P值", "校正后P值", "是否显著"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                for paragraph in table.rows[0].cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            for idx, comp in enumerate(comparisons):
                row = table.rows[idx + 1]
                row.cells[0].text = f"{comp.get('group1', '')} vs {comp.get('group2', '')}"
                row.cells[1].text = f"{comp.get('mean_diff', 0):.4f}" if comp.get("mean_diff") is not None else "-"
                p_val = comp.get("p_value")
                row.cells[2].text = _format_p_value(p_val) if p_val is not None else "-"
                adj_p = comp.get("adjusted_p_value")
                row.cells[3].text = _format_p_value(adj_p) if adj_p is not None else "-"
                row.cells[4].text = "是*" if comp.get("significant") else "否"
    
    # 矛盾结果解释
    contradiction_note = data.get("contradiction_note")
    if contradiction_note:
        doc.add_paragraph("")
        doc.add_paragraph(contradiction_note)
    
    doc.add_paragraph("* P < 0.05 表示至少有一组与其他组存在显著差异")

def _add_regression_to_doc(doc: Document, data: Dict[str, Any], field_map: Dict[str, str] = None):
    """将回归分析结果以表格形式添加到文档"""
    if field_map is None:
        field_map = {}
    
    # 模型名称
    model_name = data.get("model", "回归分析")
    p = doc.add_paragraph()
    p.add_run(f"模型: ").bold = True
    p.add_run(model_name)
    
    # 因变量与自变量说明
    target = data.get("target")
    features = data.get("features")
    if target and features:
        p = doc.add_paragraph()
        p.add_run("因变量（预测目标）: ").bold = True
        p.add_run(target)
        p = doc.add_paragraph()
        p.add_run("自变量: ").bold = True
        p.add_run(", ".join(features))
    
    # VIF共线性诊断
    vif_data = data.get("vif")
    vif_interp = data.get("vif_interpretation")
    if vif_data:
        doc.add_heading("共线性诊断 (VIF)", level=3)
        vif_table = doc.add_table(rows=len(vif_data) + 1, cols=2)
        vif_table.style = "Light Grid Accent 1"
        vif_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        header_row = vif_table.rows[0]
        header_row.cells[0].text = "变量"
        header_row.cells[1].text = "VIF"
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        for idx, (feat, vif_val) in enumerate(vif_data.items()):
            row = vif_table.rows[idx + 1]
            row.cells[0].text = feat
            row.cells[1].text = f"{vif_val:.2f}"
        
        if vif_interp:
            doc.add_paragraph(vif_interp)
    
    # R²效应量解读
    r_squared_interp = data.get("r_squared_interpretation")
    if r_squared_interp:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("模型解释力: ").bold = True
        p.add_run(r_squared_interp)
    
    # 模型指标表格
    doc.add_heading("模型指标", level=3)
    metrics_table = doc.add_table(rows=1, cols=2)
    metrics_table.style = "Light Grid Accent 1"
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 收集所有可能的指标
    metric_rows = []
    if data.get("r_squared") is not None:
        metric_rows.append(("R²", f"{data['r_squared']:.4f}"))
    if data.get("adjusted_r_squared") is not None:
        metric_rows.append(("调整后R²", f"{data['adjusted_r_squared']:.4f}"))
    f_stat = data.get("f_statistic")
    if f_stat is not None:
        metric_rows.append(("F统计量", f"{f_stat:.4f}"))
    if data.get("f_p_value") is not None:
        metric_rows.append(("F检验P值", _format_p_value(data["f_p_value"])))
    if data.get("rmse") is not None:
        metric_rows.append(("均方根误差(RMSE)", f"{data['rmse']:.4f}"))
    if data.get("accuracy") is not None:
        metric_rows.append(("准确率", f"{data['accuracy'] * 100:.2f}%"))
    if data.get("alpha") is not None:
        metric_rows.append(("正则化参数α", str(data["alpha"])))
    if data.get("l1_ratio") is not None:
        metric_rows.append(("L1比例", str(data["l1_ratio"])))
    if data.get("degree") is not None:
        metric_rows.append(("多项式次数", str(data["degree"])))
    if data.get("n_features_selected") is not None:
        metric_rows.append(("选中特征数", str(data["n_features_selected"])))
    if data.get("selected_features") is not None:
        metric_rows.append(("选中特征", ", ".join(data["selected_features"])))
    if data.get("n_positive") is not None:
        metric_rows.append(("正类样本数", str(data["n_positive"])))
    if data.get("n_negative") is not None:
        metric_rows.append(("负类样本数", str(data["n_negative"])))
    if data.get("n_samples") is not None:
        metric_rows.append(("样本量", str(data["n_samples"])))
    if data.get("direction") is not None:
        metric_rows.append(("方向", data["direction"]))
    
    # 添加表头和数据行
    header_row = metrics_table.rows[0]
    header_row.cells[0].text = "指标"
    header_row.cells[1].text = "值"
    for cell in header_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for metric_name, metric_value in metric_rows:
        row = metrics_table.add_row()
        row.cells[0].text = metric_name
        row.cells[1].text = metric_value
    
    # F统计量异常说明
    f_stat_note = data.get("f_statistic_note")
    if f_stat_note:
        doc.add_paragraph(f_stat_note)
    
    # 回归系数表格
    coefficients = data.get("coefficients")
    if coefficients:
        doc.add_heading("回归系数", level=3)
        has_p_values = data.get("p_values") is not None
        col_count = 3 if has_p_values else 2
        coef_table = doc.add_table(rows=len(coefficients) + 1, cols=col_count)
        coef_table.style = "Light Grid Accent 1"
        coef_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        header_row = coef_table.rows[0]
        header_row.cells[0].text = "变量"
        header_row.cells[1].text = "系数"
        if has_p_values:
            header_row.cells[2].text = "P值"
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # 数据行
        p_values = data.get("p_values", {})
        for idx, (name, coef) in enumerate(coefficients.items()):
            row = coef_table.rows[idx + 1]
            row.cells[0].text = name
            row.cells[1].text = f"{coef:.4f}" if isinstance(coef, (int, float)) else str(coef)
            if has_p_values and name in p_values:
                row.cells[2].text = _format_p_value(p_values[name])
    
    # 优势比表格（逻辑回归）
    odds_ratio = data.get("odds_ratio")
    if odds_ratio:
        doc.add_heading("优势比(OR值)", level=3)
        or_table = doc.add_table(rows=len(odds_ratio) + 1, cols=2)
        or_table.style = "Light Grid Accent 1"
        or_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        header_row = or_table.rows[0]
        header_row.cells[0].text = "变量"
        header_row.cells[1].text = "OR值"
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        for idx, (name, or_val) in enumerate(odds_ratio.items()):
            row = or_table.rows[idx + 1]
            row.cells[0].text = name
            row.cells[1].text = f"{or_val:.4f}"
    
    # 模型说明
    note = data.get("note")
    if note:
        doc.add_paragraph("")
        doc.add_paragraph(note)

def _add_dict_to_doc(doc: Document, data: Dict[str, Any], indent: int = 0, field_map: Dict[str, str] = None):
    if field_map is None:
        field_map = {}
    for key, value in data.items():
        display_key = field_map.get(key, key)
        if isinstance(value, dict):
            p = doc.add_paragraph()
            p.add_run(f"{display_key}:").bold = True
            if indent > 0:
                p.paragraph_format.left_indent = Inches(indent * 0.3)
            _add_dict_to_doc(doc, value, indent + 1, field_map)
        elif isinstance(value, list):
            p = doc.add_paragraph()
            p.add_run(f"{display_key}:").bold = True
            for item in value:
                if isinstance(item, dict):
                    _add_dict_to_doc(doc, item, indent + 1, field_map)
                else:
                    doc.add_paragraph(str(item), style="List Bullet")
        else:
            p = doc.add_paragraph()
            p.add_run(f"{display_key}: ").bold = True
            p.add_run(str(value))
            if indent > 0:
                p.paragraph_format.left_indent = Inches(indent * 0.3)

def create_pdf_report(
    filename: str,
    df: pd.DataFrame,
    cleaning_log: Dict[str, Any],
    analysis_results: Dict[str, Any],
    charts: Dict[str, str],
    ai_conclusion: Optional[str] = None,
    output_path: Optional[str] = None,
) -> str:
    html_content = _generate_report_html(filename, df, cleaning_log, analysis_results, charts, ai_conclusion)

    if output_path is None:
        output_path = BASE_DIR / f"report_{filename.replace('.', '_')}.pdf"

    try:
        from weasyprint import HTML
        HTML(string=html_content).write_pdf(str(output_path))
        logger.info(f"PDF report saved: {output_path}")
    except Exception as e:
        logger.error(f"WeasyPrint failed: {e}, trying pdfkit")
        try:
            import pdfkit
            pdfkit.from_string(html_content, str(output_path))
            logger.info(f"PDF report saved via pdfkit: {output_path}")
        except Exception as e2:
            logger.error(f"pdfkit also failed: {e2}")
            raise RuntimeError(f"PDF生成失败: {e2}")

    return str(output_path)

def _generate_report_html(
    filename: str,
    df: pd.DataFrame,
    cleaning_log: Dict[str, Any],
    analysis_results: Dict[str, Any],
    charts: Dict[str, str],
    ai_conclusion: Optional[str] = None,
) -> str:
    html = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body { font-family: 'SimSun', 'Times New Roman', serif; margin: 40px; line-height: 1.6; }
            h1 { color: #003366; text-align: center; border-bottom: 2px solid #003366; padding-bottom: 10px; }
            h2 { color: #003366; border-bottom: 1px solid #ccc; padding-bottom: 5px; }
            h3 { color: #333; }
            table { border-collapse: collapse; width: 100%; margin: 20px 0; }
            th, td { border: 1px solid #ddd; padding: 8px; text-align: center; }
            th { background-color: #f0f0f0; font-weight: bold; }
            img { max-width: 100%; height: auto; display: block; margin: 20px auto; }
            .conclusion { background-color: #f9f9f9; padding: 15px; border-left: 4px solid #003366; }
        </style>
    </head>
    <body>
    """

    html += f"<h1>生物医药实验数据分析报告</h1>"
    html += f"<p><strong>数据文件:</strong> {filename}</p>"
    html += f"<p><strong>数据行数:</strong> {len(df)}, <strong>列数:</strong> {len(df.columns)}</p>"

    html += "<h2>一、数据清洗摘要</h2>"
    html += f"<p>原始数据: {cleaning_log.get('original_rows', 'N/A')} 行 × {cleaning_log.get('original_cols', 'N/A')} 列</p>"
    html += f"<p>清洗后数据: {cleaning_log.get('final_rows', 'N/A')} 行 × {cleaning_log.get('final_cols', 'N/A')} 列</p>"
    
    if cleaning_log.get("steps"):
        html += "<h3>清洗步骤</h3><ul>"
        for step in cleaning_log["steps"]:
            step_name = step.get('step', '')
            step_desc = ""
            if step_name == "remove_empty_columns":
                step_desc = f"删除了 {step.get('count', 0)} 个空列"
            elif step_name == "remove_empty_rows":
                step_desc = f"删除了 {step.get('removed_count', 0)} 个空行"
            elif step_name == "remove_duplicates":
                step_desc = f"删除了 {step.get('removed_count', 0)} 个重复行"
            elif step_name == "fill_missing":
                step_desc = f"列 '{step.get('column', '')}' 使用 {step.get('strategy', '')} 填充了 {step.get('filled_count', 0)} 个缺失值"
            elif step_name == "fill_missing_categorical":
                step_desc = f"列 '{step.get('column', '')}' 使用众数填充了 {step.get('filled_count', 0)} 个缺失值"
            elif step_name == "missing_summary":
                step_desc = f"缺失值: 清洗前 {step.get('before', 0)} 个 → 清洗后 {step.get('after', 0)} 个"
            else:
                step_desc = str({k: v for k, v in step.items() if k != "step"})
            
            html += f"<li><strong>{step_name}:</strong> {step_desc}</li>"
        html += "</ul>"

    html += "<h2>二、描述性统计</h2>"
    desc_stats = analysis_results.get("descriptive_statistics", {})
    if desc_stats and "error" not in desc_stats:
        html += "<table><tr><th>变量</th><th>均值</th><th>标准差</th><th>中位数</th><th>最小值</th><th>最大值</th><th>缺失值</th></tr>"
        for col_name, stats_data in desc_stats.items():
            html += f"<tr><td>{col_name}</td><td>{stats_data.get('mean', 0):.3f}</td><td>{stats_data.get('std', 0):.3f}</td>"
            html += f"<td>{stats_data.get('median', 0):.3f}</td><td>{stats_data.get('min', 0):.3f}</td>"
            html += f"<td>{stats_data.get('max', 0):.3f}</td><td>{stats_data.get('missing', 0)}</td></tr>"
        html += "</table>"

    html += "<h2>三、图表</h2>"
    for chart_name, chart_path in charts.items():
        # Convert relative path to absolute path for proper file access
        abs_path = Path(chart_path) if Path(chart_path).is_absolute() else Path(__file__).parent.parent.parent / "charts" / chart_path
        if abs_path.exists():
            # For PDF reports, we need to embed images as base64 to ensure they show up correctly
            import base64
            try:
                with open(abs_path, "rb") as img_file:
                    img_data = base64.b64encode(img_file.read()).decode('utf-8')
                    html += f"<h3>{chart_name.replace('_', ' ').title()}</h3>"
                    html += f'<img src="data:image/png;base64,{img_data}" alt="{chart_name}">'
            except Exception as e:
                logger.error(f"Failed to embed chart {chart_name}: {e}")
                html += f"<p>[图表加载失败: {chart_name}]</p>"
        else:
            logger.error(f"Chart file not found: {abs_path}")
            html += f"<p>[图表文件不存在: {chart_path}]</p>"

    if ai_conclusion:
        html += "<h2>四、AI分析结论</h2>"
        html += f'<div class="conclusion">{ai_conclusion}</div>'

    issues = analysis_results.get("issues", [])
    if issues:
        html += "<h2>五、数据问题与建议</h2>"
        severity_order = {"高": 0, "中等": 1, "低": 2}
        sorted_issues = sorted(issues, key=lambda x: severity_order.get(x.get("severity", "低"), 3))
        for issue in sorted_issues:
            severity_icon = {"高": "🔴", "中等": "", "低": "🟢"}.get(issue.get("severity", "低"), "")
            html += f'<h3>{severity_icon} [{issue.get("severity", "")}] {issue.get("category", "")}</h3>'
            html += f'<p><strong>问题:</strong> {issue.get("problem", "")}</p>'
            html += f'<p><strong>建议:</strong> {issue.get("suggestion", "")}</p>'
            if issue.get("details"):
                html += f'<p><strong>详情:</strong> {issue.get("details", "")}</p>'

    html += "</body></html>"
    return html
